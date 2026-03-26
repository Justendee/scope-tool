"""
AI Assisted Scope Management Tool — Streamlit UI (layout only).
"""

import base64
import json
import os
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import fitz
import streamlit as st
from anthropic import Anthropic
from dotenv import load_dotenv
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from appendix_b_boilerplate import GENERAL_SCOPE_BOILERPLATE

# ── Boilerplate tag diagnostics (printed once on startup) ─────────────────────
_DIAG_ITEMS = {6: 3, 20: 17, 44: 41, 45: 42, 46: 43}  # item_number: list_index
print("\n=== BOILERPLATE TAG DIAGNOSTICS ===")
for _item_num, _idx in _DIAG_ITEMS.items():
    _raw = GENERAL_SCOPE_BOILERPLATE[_idx]
    _preview = _raw[:120].replace("\n", "\\n")
    print(f"  Item {_item_num} (index {_idx}): {_preview!r}")
print("===================================\n")
# ──────────────────────────────────────────────────────────────────────────────

SPEC_PARSE_SYSTEM_PROMPT = (
    "You are an expert construction estimator working for a general "
    "contractor. You are reading a specification division and extracting "
    "scope items for a trade subcontractor scope document."
)

SPEC_PARSE_MODEL = "claude-sonnet-4-6"

APPENDIX_B_SYSTEM_PROMPT = (
    "You are an expert construction estimator working for a "
    "general contractor in British Columbia. Your task is to "
    "generate the Specific Scope of Work section (item 56) for "
    "a subcontractor Appendix B scope document."
)

APPENDIX_B_USER_PROMPT_TEMPLATE = """TARGET TRADE: {trade_or_division}

You are generating an Appendix B specifically for the \
{trade_or_division} trade. Extract and include only scope \
items relevant to this trade from the specification and \
drawing content provided. If the uploaded documents contain \
little or no content relevant to this trade, return a minimal \
scope with a note that the relevant specification sections \
were not provided.

You have been provided with:
1. Extracted specification content for this trade
2. A drawing index showing what is visible on the drawings
3. Estimator notes for this project and this trade

RULES:
- Generate scope items specific to this project based on 
  the drawings and specifications provided. Do not generate 
  generic scope items.
- Write each scope item in plain construction language as a 
  complete standalone sentence.
- Target 15 to 30 numbered items depending on scope 
  complexity. Quality over quantity.
- Where multiple closely related items can be expressed 
  naturally in one sentence, combine them. Do not force 
  unrelated items together. If two items cover distinct 
  scope they stay as separate lines.
- Do not create heading lines with sub-lists. Every 56.x 
  item must be a complete standalone sentence.
- Consolidate detailed technical requirements into higher 
  level scope statements. For example: test all substrates 
  for moisture content using calibrated electronic moisture 
  meter becomes test all substrates for moisture content per 
  specification requirements.
- Number all items as 56.1, 56.2, 56.3 etc. Never use 57.1 
  or any other parent number.
- Do not include items already covered in the General Scope 
  of Work boilerplate such as supervision, safety, 
  submittals, warranties, Procore, and scheduling.
- If estimator notes say to exclude something, exclude it 
  even if it appears in the drawings or specifications.
- Include multiple mobilizations as the final item if the 
  project scope suggests phased work.
- Do NOT reference the project name, project type, or project 
  location in any scope item. Scope items describe the work 
  only, not what project it is for.
  Bad: 'Allow for multiple mobilizations to accommodate phased 
  construction of this library expansion project'
  Good: 'Allow for multiple mobilizations as required for the 
  duration of the project'
- Do not include technical specifications, standards, ratings, 
  measurements, product numbers, or code references in scope 
  items. The specifications and drawings are already contract 
  documents. Each scope item confirms what work is included, not 
  how it is to be performed. For example: 'Apply intumescent 
  fireproofing to exposed structural steel columns per drawings 
  and specifications' is correct. 'Apply water-based thin-film 
  intumescent fireproofing achieving 1-hour fire resistance rating 
  per CAN/ULC S101 at 130 mil dry film thickness' is too detailed 
  and must be simplified.

Before generating the 56.x scope items, output two header lines 
in this exact order:

Line 1 — CSI MasterFormat division reference:
DIVISION_REF: [division number] - [section title] / [additional 
sections if applicable]
Example: DIVISION_REF: 09 91 00 - Painting / 09 96 46 - 
Intumescent Painting

Line 2 - INTRO: Write a short phrase describing only the 
primary work for this trade. Maximum 12 words. No verbs like 
'supply and'. Start directly with the type of work.
Do NOT include 'as shown on drawings and specifications' or 
'in accordance with drawings and specifications' in the phrase 
— those are added automatically by the document generator.
Do NOT include any reference to drawings, specifications, or 
contract documents in the phrase.
Do NOT end the phrase with a comma or period. The phrase must 
end with the last word of the work description only.
Do NOT include the project name, project location, or any 
project context in the INTRO phrase. Describe the type of work 
only, not where or for what project it is being done.
Good example for demo: 
'demolition and abatement works'
Bad example (contains project context): 
'demolition and abatement work for library expansion and 
renovation'
Good example for painting: 
'field-applied painting, including preparation, priming and 
finishing of all surfaces'
Bad example (contains project name): 
'interior painting of the VPL Marpole Branch Library'
Bad example (too long): 
'all interior painting and intumescent fireproofing, including 
surface preparation, priming, finish coats, and repaints of 
existing surfaces'
Bad example (references drawings): 
'demolition and abatement works per drawings and specifications'
Bad example (references specifications): 
'field-applied painting, including preparation, priming and 
finishing of all surfaces per specifications'
Bad example (contains boilerplate): 
'field-applied painting as shown on drawings and specifications'
Format: INTRO: [short phrase only, no full sentence]

Then generate the 56.x scope items on subsequent lines.

ESTIMATOR NOTES FOR THIS PROJECT: {project_notes}
ESTIMATOR NOTES FOR THIS TRADE: {division_notes}

SPECIFICATION CONTENT FOR THIS TRADE:
{spec_text}

DRAWING INDEX CONTENT RELEVANT TO THIS TRADE:
{drawing_index_content}

Output format (strict):
- Line 1: DIVISION_REF line
- Line 2: INTRO line
- Remaining lines: numbered 56.x scope items only
Do not include any other text, headings, or explanation."""

INDEX_DRAWINGS_SYSTEM_PROMPT = (
    "You are an expert construction estimator reviewing project "
    "drawings. Extract structured information from each drawing sheet."
)

INDEX_DRAWINGS_USER_PROMPT = (
    "For each drawing sheet in this batch, identify and return in "
    "JSON format:\n"
    "- sheet_number\n"
    "- discipline\n"
    "- drawing_title\n"
    "- trades_referenced\n"
    "- scope_notes (key scope items visible on this sheet)\n"
    "- cross_references (any other documents or drawings referenced)\n"
    "Return a JSON array with one object per sheet."
)

PROJECTS_DIR = Path(__file__).resolve().parent / "projects"
PROJECTS_DIR.mkdir(exist_ok=True)

# Fallback paths for when no project is active (backward-compat)
_APP_DIR = Path(__file__).resolve().parent


def _get_project_folder(project_number: str) -> Path:
    """Return the folder path for project_number, or the app root.

    This function is PURE PATH COMPUTATION — it never creates the folder.
    Call folder.mkdir(parents=True, exist_ok=True) explicitly at the point
    where files are actually being written.
    """
    safe = (project_number or "").strip()
    if not safe:
        return _APP_DIR
    safe = re.sub(r'[<>:"/\\|?*\s]+', "_", safe).strip("_")
    if not safe:
        return _APP_DIR
    return PROJECTS_DIR / safe

DRAWING_PAGE_RENDER_DPI = 150
INDEX_BATCH_SIZE = 3


def _sanitize_filename(value: str) -> str:
    """Make a Windows-safe filename segment (spaces allowed)."""
    value = value.strip()
    if not value:
        return "Unknown"
    # Replace invalid Windows filename characters with spaces
    value = re.sub(r'[<>:"/\\|?*]+', " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def _project_field(value: str) -> str:
    """Use sidebar value or Unknown for filenames."""
    v = (value or "").strip()
    return v if v else "Unknown"


def _escape_braces_for_format(s: str) -> str:
    """Allow str.format on templates when user/model text contains { or }."""
    return s.replace("{", "{{").replace("}", "}}")


def _spaced_download_filename(parts: list[str], ext: str) -> str:
    """Build names like: Part1 - Part2 - Part3.docx (spaces, no underscores)."""
    safe = [_sanitize_filename(p) for p in parts]
    return " - ".join(safe) + f".{ext}"


def _parse_appendix_b_response(api_text: str) -> tuple[str, str, str]:
    """
    Parse Claude's Appendix B response into three parts.

    Returns:
        (division_ref, intro_text, scope_body)
    Expects lines starting with DIVISION_REF: and INTRO: before the
    56.x items; falls back gracefully if either header is missing.
    """
    text = (api_text or "").strip()
    if not text:
        return "", "", ""

    lines = text.splitlines()
    division_ref = ""
    intro_text = ""
    body_start = 0

    for idx, line in enumerate(lines):
        stripped = line.strip()
        upper = stripped.upper()
        if upper.startswith("DIVISION_REF:"):
            division_ref = stripped.split(":", 1)[1].strip()
            body_start = idx + 1
        elif upper.startswith("INTRO:"):
            intro_text = stripped[len("INTRO:"):].strip()
            body_start = idx + 1
        elif stripped.startswith("56."):
            body_start = idx
            break

    scope_body = "\n".join(lines[body_start:]).strip()
    return division_ref, intro_text, scope_body


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float, bool)):
        return str(value)
    # For lists/dicts returned by the model, store them as readable JSON.
    try:
        return json.dumps(value, ensure_ascii=False)
    except Exception:
        return str(value)


def _scope_items_to_docx_bytes(scope_items_text: str) -> bytes:
    doc = Document()
    doc.add_heading("Scope Summary", level=1)
    doc.add_paragraph("Extracted Scope Items")

    for line in scope_items_text.splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue
        doc.add_paragraph(cleaned, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _appendix_b_word_bytes(
    project_number: str,
    trade_or_division: str,
    division_reference: str,
    specific_scope_items: str,
    entity_name: str,
    intro_text: str = "",
    subcontractor_name: str = "",
) -> bytes:
    """Full Appendix B Word document – TNR 11 pt, 0.75-inch margins, hanging indents."""
    num = _project_field(project_number)
    tname = _project_field(trade_or_division)
    raw_div = (division_reference or "").strip()
    # CSI numbers only (for Division line): strip section titles after each hyphen
    if raw_div:
        _parts = re.split(r"\s*/\s*", raw_div)
        _nums_only = [re.sub(r"\s*[-\u2013]\s*.*", "", p).strip() for p in _parts]
        div_numbers = " / ".join(n for n in _nums_only if n)
    else:
        div_numbers = ""
    div_display = div_numbers or "[Division reference pending]"
    # Full reference with em dashes (for item 1.2)
    div_full_emdash = raw_div.replace(" - ", " \u2013 ") if raw_div else "[Division reference pending]"
    ent = _project_field(entity_name)
    sub_display = (subcontractor_name or "").strip() or "TBD"

    doc = Document()

    # ── Margins (0.75 inch all sides) ────────────────────────────────────────
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(0.75))

    # ── Page header (every page) ──────────────────────────────────────────────
    header = sec.header
    for hp in header.paragraphs:
        hp.clear()
    hp1 = header.paragraphs[0]
    hp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh1 = hp1.add_run("APPENDIX B \u2013 SCOPE OF WORK")
    rh1.font.name = "Times New Roman"
    rh1.font.size = Pt(10)
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh2 = hp2.add_run(f"Contract#{num} \u2013 Job#{num}")
    rh2.font.name = "Times New Roman"
    rh2.font.size = Pt(10)

    # ── Page footer (every page) ──────────────────────────────────────────────
    # Single paragraph; tab stops: center at 3.5" (5040 twips) and right at 7"
    # (10080 twips) relative to the left margin of a 7"-wide text area.
    # Left: label  |  Center: date  |  Right: Page X of Y
    footer = sec.footer
    for fp in footer.paragraphs:
        fp.clear()
    fp = footer.paragraphs[0]
    pPr = fp._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    for pos_twips, tab_type in ((5040, "center"), (10080, "right")):
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), tab_type)
        tab.set(qn("w:pos"), str(pos_twips))
        tabs_el.append(tab)
    pPr.append(tabs_el)

    def _footer_run(para, text: str) -> None:
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    def _field_run(para, field_name: str) -> None:
        """Insert a simple Word field (PAGE or NUMPAGES) as a run."""
        r = para.add_run()
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        for ftype, itext in (("begin", None), (None, field_name), ("end", None)):
            if ftype is not None:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), ftype)
                r._r.append(fc)
            else:
                it = OxmlElement("w:instrText")
                it.set(qn("xml:space"), "preserve")
                it.text = f" {itext} "
                r._r.append(it)

    _footer_run(fp, "SSP - CCA 1 -2008 Contract and SC\u2019s")
    _footer_run(fp, "\t")
    _footer_run(fp, "Updated Feb 7, 2022")
    _footer_run(fp, "\t")
    _footer_run(fp, "Page ")
    _field_run(fp, "PAGE")
    _footer_run(fp, " of ")
    _field_run(fp, "NUMPAGES")

    # ── Default body style: TNR 11pt, 6pt space after, no space before ────────
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    # ── Paragraph helpers ─────────────────────────────────────────────────────
    # Numbered items: hanging indent so text wraps under the first word.
    #   first line starts at left margin (0"), text wraps at 0.5".
    #   Use "\t" between number and body text.
    # Sub-items (1.1, 56.x): first line starts at 0.5", text wraps at 1.0".
    #   Use "\t" between sub-number and body text.

    # Matches UNDERLINE:...:/UNDERLINE  and  [ITALIC]...[/ITALIC]
    _inline_markup_pat = re.compile(
        r"UNDERLINE:(.*?):/UNDERLINE|\[ITALIC\](.*?)\[/ITALIC\]", re.DOTALL
    )

    def _add_inline_markup_runs(
        para,
        text: str,
        *,
        bold: bool = False,
        strike: bool = False,
        italic: bool = False,
    ) -> None:
        """Add runs, rendering UNDERLINE:...:/UNDERLINE and [ITALIC]...[/ITALIC] spans."""
        last = 0
        for m in _inline_markup_pat.finditer(text):
            pre = text[last : m.start()]
            if pre:
                r = para.add_run(pre)
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = bold
                r.font.strike = strike
                r.italic = italic
            is_underline = m.group(1) is not None
            span_text = m.group(1) if is_underline else m.group(2)
            rs = para.add_run(span_text)
            rs.font.name = "Times New Roman"
            rs.font.size = Pt(11)
            rs.bold = bold
            rs.font.strike = strike
            rs.italic = italic if is_underline else True
            if is_underline:
                rs.font.underline = WD_UNDERLINE.SINGLE
            last = m.end()
        tail = text[last:]
        if tail:
            r = para.add_run(tail)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.bold = bold
            r.font.strike = strike
            r.italic = italic

    def _para(
        text: str,
        *,
        bold: bool = False,
        underline: bool = False,
        strike: bool = False,
        size: int = 11,
        center: bool = False,
        hanging: bool = False,
        sub_item: bool = False,
        plain_indent: float = 0.0,
    ):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if hanging:
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
        elif sub_item:
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.first_line_indent = Inches(-0.5)
        elif plain_indent:
            p.paragraph_format.left_indent = Inches(plain_indent)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.font.underline = WD_UNDERLINE.SINGLE if underline else False
        r.font.strike = strike
        return p

    def _heading(text: str) -> None:
        _para(text, bold=True, underline=True)

    def _numbered(
        num_str: str,
        body: str,
        *,
        bold: bool = False,
        strike: bool = False,
        italic: bool = False,
    ) -> None:
        """Hanging-indent numbered item: number\tbody."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{num_str}\t")
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(11)
        rn.bold = bold
        rn.font.strike = strike
        rn.italic = italic
        _add_inline_markup_runs(p, body, bold=bold, strike=strike, italic=italic)

    def _sub_item(num_str: str, body: str) -> None:
        """Sub-item: starts at 0.5", text wraps at 1.0"."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.0)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{num_str}\t")
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(11)
        rb = p.add_run(body)
        rb.font.name = "Times New Roman"
        rb.font.size = Pt(11)

    # ── Title: centered, TNR 12pt, underlined ─────────────────────────────────
    _para(f"Contract # {num}-??? \u2013 {sub_display}", underline=True, size=12, center=True)

    # ── Preamble & division reference ─────────────────────────────────────────
    _para(
        "In addition to the scope of work indicated in the Contract Documents, "
        "the Contract Scope of Work includes but is not limited to:",
        bold=True,
    )
    _para(f"Division \u2013 {div_display}", bold=True)

    # ── Item 1 ────────────────────────────────────────────────────────────────
    _numbered(
        "1.",
        "The Scope of Work of this Contract includes but is not limited to, "
        "the supply of labour, materials, equipment and accessories required to "
        "execute and complete the work as per the Contract Documents, IFC "
        "Drawings and Specifications, including but not limited to the "
        "following sections:",
    )
    _sub_item("1.1", "Division 1 \u2013 General Requirements")
    _sub_item("1.2", f"Section {div_full_emdash}")
    _sub_item("1.3", "Related Sections in Specifications")

    # ── Item 2 heading ────────────────────────────────────────────────────────
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.first_line_indent = Inches(-0.5)
    p2.paragraph_format.space_after = Pt(4)
    r2n = p2.add_run("2.\t")
    r2n.font.name = "Times New Roman"
    r2n.font.size = Pt(11)
    r2h = p2.add_run("General Scope of Work:")
    r2h.font.name = "Times New Roman"
    r2h.font.size = Pt(11)
    r2h.font.underline = WD_UNDERLINE.SINGLE

    # ── Items 3–55 ────────────────────────────────────────────────────────────
    for i, raw in enumerate(GENERAL_SCOPE_BOILERPLATE):
        item_num = i + 3
        is_strike = raw.startswith("STRIKETHROUGH:")
        body = raw[len("STRIKETHROUGH:"):].strip() if is_strike else raw
        if item_num == 7:
            body = body.replace("Scott Special Projects Ltd.", ent)
        _numbered(f"{item_num}.", body, strike=is_strike, italic=is_strike)

    # ── Item 56 heading (standalone bold + underlined, no number) ────────────
    _heading("Specific Scope of Work:")

    # ── Item 56 intro sentence (numbered) ─────────────────────────────────────
    _verb = DIVISION_VERB_PHRASES.get(trade_or_division, "supply and install")
    _phrase = (intro_text or "").strip().rstrip(",.") or f"{tname} scope"
    if _verb == "supply and application of":
        _connector = f"for the {_verb}"
    else:
        _connector = f"to {_verb}"
    _intro = (
        f"Provide all labour, materials and equipment {_connector} "
        f"{_phrase}, in accordance with drawings and specifications, "
        f"including, but not limited to:"
    )
    _numbered("56.", _intro)

    # ── 56.x sub-items ───────────────────────────────────────────────────────
    _sub_pat = re.compile(r"^(56\.\d+)\s+(.*)")
    for line in specific_scope_items.splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue
        m = _sub_pat.match(cleaned)
        if m:
            _sub_item(m.group(1), m.group(2))
        else:
            _para(cleaned, plain_indent=1.0)

    # ── Item 57 ───────────────────────────────────────────────────────────────
    _numbered(
        "57.",
        "Provide own supervision and first-aid coverage necessary to meet "
        "WorkSafeBC requirements for this trade\u2019s scope in case of work "
        "independent of GC supervision",
    )

    # ── Scope of Work Exclusions ──────────────────────────────────────────────
    _heading("Scope of Work Exclusions:")
    _para("To be completed upon receipt of trade quotes.")

    # ── Scheduling of Work ────────────────────────────────────────────────────
    _heading("Scheduling of Work:")
    _para("As per attached Appendix D Project Schedule")

    # ── Contract Administration ───────────────────────────────────────────────
    _heading("Contract Administration")
    ca_items = [
        (
            "58.",
            "Progress claim summary including base contract breakdown to be "
            "submitted with pre-approved schedule of values in monthly progress "
            "invoices",
        ),
        (
            "59.",
            f"Certificate of Liability Insurance document c/w {ent} as "
            "additional specific named insured as per SSP Insurance template",
        ),
        (
            "60.",
            "Provide complete listing of all charge out rates for the durations "
            "of the work of the contract. Rates include all "
            "labour/material/equipment costs, escalation, overhead, and fee "
            "mark-ups",
        ),
        (
            "61.",
            "Provision of WSBC clearance letter prior to release of first "
            "progress payment, and statutory declaration prior to release of "
            "second progress payment",
        ),
        (
            "62.",
            "Final payment may be withheld until all required closeout "
            "documentation has been submitted",
        ),
    ]
    for ca_num, ca_body in ca_items:
        _numbered(ca_num, ca_body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _drawing_index_to_xlsx_bytes(drawing_index: list[dict]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Drawing Index"

    headers = [
        "Sheet Number",
        "Title",
        "Discipline",
        "Trades Referenced",
        "Scope Notes",
        "Cross References",
    ]
    ws.append(headers)

    header_font = Font(bold=True)
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for item in drawing_index:
        ws.append(
            [
                _stringify_cell(item.get("sheet_number", "")),
                _stringify_cell(item.get("drawing_title", "")),
                _stringify_cell(item.get("discipline", "")),
                _stringify_cell(item.get("trades_referenced", "")),
                _stringify_cell(item.get("scope_notes", "")),
                _stringify_cell(item.get("cross_references", "")),
            ]
        )

    # Basic column width sizing
    for col_idx, header in enumerate(headers, start=1):
        max_len = len(header)
        for row_idx in range(2, ws.max_row + 1):
            cell_val = ws.cell(row=row_idx, column=col_idx).value
            if cell_val is None:
                continue
            max_len = max(max_len, len(str(cell_val)))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 60)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def load_env_from_dotenv() -> str | None:
    """Load `.env` and return `ANTHROPIC_API_KEY` from the environment."""
    load_dotenv()
    return os.environ.get("ANTHROPIC_API_KEY")


ANTHROPIC_API_KEY = load_env_from_dotenv()

def _restore_project_files(project_folder: Path) -> None:
    """Load scope_summary.txt and drawing_index.json from folder into session state."""
    if "extracted_scope_items_text" not in st.session_state:
        _ss = project_folder / "scope_summary.txt"
        if _ss.exists():
            try:
                _saved = _ss.read_text(encoding="utf-8").strip()
                if _saved:
                    st.session_state["extracted_scope_items_text"] = _saved
                    st.session_state["scope_summary_ready"] = True
            except OSError:
                pass
    _di = project_folder / "drawing_index.json"
    if _di.exists():
        try:
            _data = json.loads(_di.read_text(encoding="utf-8"))
            if isinstance(_data, list) and _data:
                st.session_state["drawing_index_ready"] = True
        except Exception:
            pass


def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract plain text from a PDF using PyMuPDF, page by page."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text())
        return "\n".join(parts)
    finally:
        doc.close()


def _page_to_png_highres(page: fitz.Page) -> bytes:
    """Render a PDF page to a high-resolution PNG using PyMuPDF."""
    zoom = DRAWING_PAGE_RENDER_DPI / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    w, h = pix.width, pix.height
    if w > 8000 or h > 8000:
        scale = min(7500 / w, 7500 / h)
        mat = fitz.Matrix(zoom * scale, zoom * scale)
        pix = page.get_pixmap(matrix=mat, alpha=False)
    return pix.tobytes("png")


def _parse_json_array_from_model_text(text: str) -> list:
    """Parse a JSON array from Claude output, tolerating markdown fences."""
    raw = text.strip()
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", raw)
    if fence:
        raw = fence.group(1).strip()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        start = raw.find("[")
        end = raw.rfind("]")
        if start == -1 or end == -1 or end <= start:
            raise
        data = json.loads(raw[start : end + 1])
    if not isinstance(data, list):
        raise ValueError("Model response was not a JSON array")
    return data


def _build_index_batch_content(png_pages: list[bytes]) -> list[dict]:
    """Anthropic Messages API content blocks: ordered images then the user prompt."""
    blocks: list[dict] = []
    for png in png_pages:
        b64 = base64.standard_b64encode(png).decode("utf-8")
        blocks.append(
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": b64,
                },
            }
        )
    blocks.append({"type": "text", "text": INDEX_DRAWINGS_USER_PROMPT})
    return blocks


def index_drawings(
    pdf_bytes: bytes,
    project_number: str,
    project_name: str,
    project_folder: Path | None = None,
) -> None:
    """
    Render each page as a high-res image, index in batches of three via Claude,
    merge into `drawing_index.json` in project_folder, and show a summary.
    """
    if project_folder is None:
        project_folder = _get_project_folder(project_number)
    # Ensure folder exists before writing drawing_index.json
    if project_folder != _APP_DIR:
        project_folder.mkdir(parents=True, exist_ok=True)
    if not ANTHROPIC_API_KEY:
        st.error(
            "ANTHROPIC_API_KEY is not set. Add it to your `.env` file in the "
            "project folder."
        )
        return

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        n_pages = doc.page_count
        if n_pages == 0:
            st.warning("The PDF has no pages to index.")
            return

        with st.spinner("Rendering drawing pages at high resolution…"):
            page_pngs: list[bytes] = []
            for page in doc:
                page_pngs.append(_page_to_png_highres(page))
    finally:
        doc.close()

    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    combined: list[dict] = []
    n_batches = (n_pages + INDEX_BATCH_SIZE - 1) // INDEX_BATCH_SIZE
    progress_slot = st.empty()
    progress_slot.progress(0, text="Indexing drawings…")

    try:
        for b in range(n_batches):
            start = b * INDEX_BATCH_SIZE
            batch_pngs = page_pngs[start : start + INDEX_BATCH_SIZE]
            content = _build_index_batch_content(batch_pngs)
            with st.spinner(
                f"Indexing batch {b + 1} of {n_batches} ({len(batch_pngs)} sheet(s))…"
            ):
                message = client.messages.create(
                    model=SPEC_PARSE_MODEL,
                    max_tokens=8192,
                    system=INDEX_DRAWINGS_SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": content}],
                )
            reply_parts: list[str] = []
            for block in message.content:
                if block.type == "text":
                    reply_parts.append(block.text)
            reply_text = "".join(reply_parts)
            batch_rows = _parse_json_array_from_model_text(reply_text)
            for row in batch_rows:
                if isinstance(row, dict):
                    combined.append(row)
            progress_slot.progress((b + 1) / n_batches, text="Indexing drawings…")
    except Exception as exc:
        progress_slot.empty()
        st.error(f"Drawing index failed: {exc}")
        return

    progress_slot.empty()

    _drawing_index_path = project_folder / "drawing_index.json"
    try:
        with open(_drawing_index_path, "w", encoding="utf-8") as f:
            json.dump(combined, f, indent=2, ensure_ascii=False)
    except OSError as exc:
        st.error(f"Could not write {_drawing_index_path}: {exc}")
        return

    st.session_state["drawing_index_ready"] = True

    st.success(f"Saved drawing index to `{_drawing_index_path.name}`.")

    st.subheader("Drawing index summary")
    st.write(
        f"**Sheets indexed:** {len(combined)} row(s) extracted from "
        f"{n_pages} drawing page(s), sent in {n_batches} batch(es) of up to "
        f"{INDEX_BATCH_SIZE} page(s) each."
    )
    table_rows = [
        {
            "Sheet number": r.get("sheet_number", ""),
            "Title": r.get("drawing_title", ""),
        }
        for r in combined
    ]
    st.dataframe(table_rows, use_container_width=True, hide_index=True)

    try:
        drawing_index_data = json.loads(_drawing_index_path.read_text(encoding="utf-8"))
    except Exception as exc:
        st.error(f"Could not load `{_drawing_index_path.name}` for Excel export: {exc}")
        return

    xlsx_bytes = _drawing_index_to_xlsx_bytes(drawing_index_data)
    st.download_button(
        label="Download Drawing Index (.xlsx)",
        data=xlsx_bytes,
        file_name=_spaced_download_filename(
            [
                _project_field(project_number),
                _project_field(project_name),
                "Drawing Index",
            ],
            "xlsx",
        ),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def parse_spec_division(spec_text: str, project_folder: Path | None = None) -> None:
    """
    Send extracted specification text to Claude and show scope items in the
    main panel. Saves scope_summary.txt to project_folder.
    """
    if not ANTHROPIC_API_KEY:
        st.error(
            "ANTHROPIC_API_KEY is not set. Add it to your `.env` file in the "
            "project folder."
        )
        return
    stripped = spec_text.strip()
    if not stripped:
        st.warning("No specification text to parse. Upload a PDF with text.")
        return

    user_prompt = (
        "Read this specification division and extract all scope items that "
        "would be required for a subcontractor performing this work. List "
        "each scope item on a separate line. Be specific and project-focused. "
        "Do not include administrative or submittal requirements. "
        f"Specification text: {stripped}"
    )

    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    try:
        with st.spinner("Calling Claude to extract scope items…"):
            message = client.messages.create(
                model=SPEC_PARSE_MODEL,
                max_tokens=8192,
                system=SPEC_PARSE_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_prompt}],
            )
    except Exception as exc:
        st.error(f"Anthropic API request failed: {exc}")
        return

    reply_parts: list[str] = []
    for block in message.content:
        if block.type == "text":
            reply_parts.append(block.text)
    result = "".join(reply_parts)

    st.session_state["extracted_scope_items_text"] = result
    st.session_state["scope_summary_ready"] = True

    _scope_path = (project_folder or _APP_DIR) / "scope_summary.txt"
    try:
        # Ensure folder exists before writing (folder creation is intentional here)
        if project_folder and project_folder != _APP_DIR:
            project_folder.mkdir(parents=True, exist_ok=True)
        _scope_path.write_text(result, encoding="utf-8")
    except OSError as exc:
        st.warning(f"Could not save scope summary to disk: {exc}")

    st.subheader("Extracted Scope Items")
    st.markdown(result)

    return result


def _has_scope_summary_output() -> bool:
    t = st.session_state.get("extracted_scope_items_text")
    return bool(t and str(t).strip())


def _has_drawing_index_file(project_folder: Path | None = None) -> bool:
    path = (project_folder or _APP_DIR) / "drawing_index.json"
    if not path.exists():
        return False
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return isinstance(data, list) and len(data) > 0
    except Exception:
        return False


def generate_appendix_b(
    project_notes: str,
    division_notes: str,
    trade_or_division: str,
    project_folder: Path | None = None,
) -> None:
    """
    Build Appendix B Item 56 scope from scope summary, drawing index,
    and estimator notes via Claude.
    """
    if project_folder is None:
        project_folder = _APP_DIR
    if not ANTHROPIC_API_KEY:
        st.error(
            "ANTHROPIC_API_KEY is not set. Add it to your `.env` file in the "
            "project folder."
        )
        return

    spec_text = (st.session_state.get("extracted_scope_items_text") or "").strip()
    if not spec_text:
        st.warning("No extracted specification scope text found. Run Generate Scope Summary first.")
        return

    _di_path = project_folder / "drawing_index.json"
    try:
        drawing_data = json.loads(_di_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        st.warning("drawing_index.json not found. Run Generate Scope first.")
        return
    except Exception as exc:
        st.error(f"Could not read drawing index: {exc}")
        return

    if not isinstance(drawing_data, list) or len(drawing_data) == 0:
        st.warning("Drawing index is empty. Run Index Drawings first.")
        return

    drawing_index_content = json.dumps(drawing_data, indent=2, ensure_ascii=False)

    dn = division_notes.strip() if division_notes else ""
    trade_line = f"Trade or division: {trade_or_division}"
    division_notes_payload = f"{trade_line}\n\n{dn}" if dn else trade_line

    user_prompt = APPENDIX_B_USER_PROMPT_TEMPLATE.format(
        trade_or_division=_escape_braces_for_format(trade_or_division),
        project_notes=_escape_braces_for_format(
            project_notes.strip() if project_notes else ""
        ),
        division_notes=_escape_braces_for_format(division_notes_payload),
        spec_text=_escape_braces_for_format(spec_text),
        drawing_index_content=_escape_braces_for_format(drawing_index_content),
    )

    st.session_state.pop("appendix_b_scope_items", None)
    st.session_state.pop("division_ref", None)
    st.session_state.pop("appendix_b_intro", None)

    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    try:
        with st.spinner("Generating Appendix B Specific Scope of Work…"):
            message = client.messages.create(
                model=SPEC_PARSE_MODEL,
                max_tokens=8192,
                system=APPENDIX_B_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_prompt}],
            )
    except Exception as exc:
        st.error(f"Anthropic API request failed: {exc}")
        return

    reply_parts: list[str] = []
    for block in message.content:
        if block.type == "text":
            reply_parts.append(block.text)
    result = "".join(reply_parts)

    division_ref, intro_text, scope_body = _parse_appendix_b_response(result)
    st.session_state["division_ref"] = division_ref
    st.session_state["appendix_b_intro"] = intro_text
    st.session_state["appendix_b_scope_items"] = scope_body

    st.subheader("Appendix B - Specific Scope of Work (Item 56)")
    st.markdown(scope_body)


ENTITY_OPTIONS = [
    "SCOTT Construction Ltd",
    "SCOTT Special Projects Ltd",
    "Scott DB Services Ltd",
    "SCOTT Construction Management Ltd",
    "SCOTT Construction Ontario Inc",
]

PROJECT_TYPE_OPTIONS = [
    "Commercial",
    "Residential",
    "Industrial",
    "Special Projects",
    "COV",
]

DIVISION_TABS = [
    "2. Demo and Abatement",
    "2. Demolition",
    "2. Excavating",
    "2. Landscaping",
    "3. Concrete Works",
    "4. Masonry",
    "5. Steel",
    "6. Framing",
    "6. Millwork",
    "7. Roofing",
    "7. Cladding",
    "8. Doors",
    "8. Glazing",
    "9. GWB",
    "9. Flooring",
    "9. Paint",
    "14. Conveying Systems",
    "15. Mechanical",
    "16. Electrical",
]

DIVISION_VERB_PHRASES: dict[str, str] = {
    "2. Demo and Abatement": "complete",
    "2. Demolition": "complete",
    "2. Excavating": "complete",
    "2. Landscaping": "complete",
    "3. Concrete Works": "complete",
    "4. Masonry": "supply and install",
    "5. Steel": "supply and install",
    "6. Framing": "supply and install",
    "6. Millwork": "supply and install",
    "7. Roofing": "supply and install",
    "7. Cladding": "supply and install",
    "8. Doors": "supply and install",
    "8. Glazing": "supply and install",
    "9. GWB": "supply and install",
    "9. Flooring": "supply and install",
    "9. Paint": "supply and application of",
    "14. Conveying Systems": "supply and install",
    "15. Mechanical": "supply and install",
    "16. Electrical": "supply and install",
}


def _exclusive_checkboxes(prefix: str, options: list[str]) -> str:
    """
    Render mutually exclusive checkboxes.
    Mutual exclusion is enforced via on_change callbacks, which fire at the
    START of the next script run before any widgets are rendered, making it
    safe to write to other widget-bound keys from within them.
    """
    for opt in options:
        def _on_change(selected=opt):
            if st.session_state.get(f"{prefix}_{selected}"):
                # Just checked: uncheck all others, record value
                for o in options:
                    if o != selected:
                        st.session_state[f"{prefix}_{o}"] = False
                st.session_state[f"{prefix}_value"] = selected
            else:
                # Just unchecked: if nothing else is checked, restore it
                if not any(
                    st.session_state.get(f"{prefix}_{o}") for o in options if o != selected
                ):
                    st.session_state[f"{prefix}_{selected}"] = True
        st.checkbox(opt, key=f"{prefix}_{opt}", on_change=_on_change)
    return st.session_state.get(f"{prefix}_value", options[0])


# ── File-type labels for the uploader ────────────────────────────────────────
FILE_TYPE_OPTIONS = ["Drawings", "Specifications", "Statement of Work", "Hazmat Report", "Other"]

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Scope Management Tool",
    page_icon="📐",
    layout="wide",
)

st.title("AI Assisted Scope Management")

# ── One-time checkbox state initialization (must run before widgets render) ──
# Populates the keys used by st.checkbox in _exclusive_checkboxes so that
# Streamlit sees them already present before the widgets are instantiated.
for _o in ENTITY_OPTIONS:
    _k = f"entity_{_o}"
    if _k not in st.session_state:
        st.session_state[_k] = (_o == ENTITY_OPTIONS[0])
if "entity_value" not in st.session_state:
    st.session_state["entity_value"] = ENTITY_OPTIONS[0]

for _o in PROJECT_TYPE_OPTIONS:
    _k = f"ptype_{_o}"
    if _k not in st.session_state:
        st.session_state[_k] = (_o == PROJECT_TYPE_OPTIONS[0])
if "ptype_value" not in st.session_state:
    st.session_state["ptype_value"] = PROJECT_TYPE_OPTIONS[0]

# ── Apply staged widget values (must run before widgets are instantiated) ────
# The Load button cannot write to text_input keys after those widgets render,
# so it stores values in staging keys. We consume them here on the next run.
if "_staged_project_number" in st.session_state:
    st.session_state["project_number_input"] = st.session_state.pop("_staged_project_number")
if "_staged_project_name" in st.session_state:
    st.session_state["project_name_input"] = st.session_state.pop("_staged_project_name")

# ── Reset load selectbox after deletion ──────────────────────────────────────
# _staged_load_select is set to None by the _reset_project handler after a
# deletion. Consuming it here (before the selectbox widget renders) ensures
# the selectbox defaults to "— select —" even if the browser tries to restore
# the previously selected (now-deleted) project name.
if "_staged_load_select" in st.session_state:
    st.session_state.pop("_staged_load_select")
    st.session_state.pop("load_project_select", None)  # force selectbox to default


# ── Process deferred state-clear requests (must run before any widgets) ──────
# New Project / Delete Project set "_reset_project" before st.rerun() so that
# widget-bound keys can be cleared here, before those widgets are rendered.
if st.session_state.pop("_reset_project", False):
    import shutil as _shutil
    _delete_folder = st.session_state.pop("_delete_project_folder", None)
    print(f"[RESET] _reset_project fired. _delete_folder = {_delete_folder!r}")
    if _delete_folder:
        _df_path = Path(_delete_folder)
        print(f"[RESET] About to delete: {_delete_folder}")
        print(f"[RESET] Folder exists before deletion: {_df_path.exists()}")
        if _df_path.exists():
            try:
                _shutil.rmtree(_delete_folder)
                print(f"[RESET] shutil.rmtree completed")
                print(f"[RESET] Folder still exists after deletion: {_df_path.exists()}")
                # Record the deleted folder name so the project_info.json save
                # block can refuse to recreate it, and force the load selectbox
                # to reset (guards against browser-side widget state restoration).
                _deleted_name = _df_path.name
                st.session_state.setdefault("_recently_deleted", set()).add(_deleted_name)
                st.session_state["_staged_load_select"] = None
                print(f"[RESET] Recorded {_deleted_name!r} in _recently_deleted (permanent for session)")
            except Exception as _e:
                st.error(f"Could not delete project folder: {_e}")
                print(f"[RESET] rmtree FAILED: {_e}")
        else:
            print(f"[RESET] Path does not exist — rmtree skipped")
    else:
        # New Project (no deletion): clear the recently-deleted guard so the
        # user can re-create a project with the same number if they choose to.
        st.session_state.pop("_recently_deleted", None)
        print("[RESET] _delete_folder is None — New Project flow; _recently_deleted cleared")
    # NOTE: entity_* and ptype_* keys are intentionally excluded here.
    # Those keys are widget-bound and managed by _exclusive_checkboxes.
    # Clearing them would cause a StreamlitAPIException if their checkboxes
    # are rendered in the same run.
    for _k in [
        "project_number_input", "project_name_input",
        "extracted_scope_items_text", "scope_summary_ready",
        "drawing_index_ready", "appendix_b_scope_items",
        "division_ref", "appendix_b_intro", "file_labels",
        "load_project_select", "_confirm_delete_input",
        "_show_delete_confirm",
    ]:
        st.session_state.pop(_k, None)
    # Rerun once more so the sidebar re-renders from scratch with a fresh
    # disk scan. Without this extra rerun the selectbox options list may
    # still reflect the pre-deletion state because Streamlit commits widget
    # trees within a single run before the cleaned session state propagates.
    st.rerun()

# ── Helper: list existing project folders ────────────────────────────────────
def _list_projects() -> list[str]:
    if not PROJECTS_DIR.exists():
        print("[LIST_PROJECTS] projects/ dir does not exist — returning []")
        return []
    result = sorted(
        d.name for d in PROJECTS_DIR.iterdir()
        if d.is_dir() and not d.name.startswith(".")
    )
    print(f"[LIST_PROJECTS] Scanned disk — found: {result}")
    return result


# ═══════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.header("Project")

    # ── Project number / name ────────────────────────────────────────────────
    project_number = st.text_input(
        "Project Number",
        key="project_number_input",
        placeholder="e.g. 5246",
    )
    project_name = st.text_input(
        "Project Name",
        key="project_name_input",
        placeholder="e.g. Marpole Library Expansion",
    )

    # Compute active project folder and restore files if project number is set
    active_project_folder = _get_project_folder(project_number)
    if project_number:
        _restore_project_files(active_project_folder)

        # Persist project name to disk so it can be loaded later
        _pinfo = active_project_folder / "project_info.json"
        _pname_to_save = project_name.strip()
        try:
            _current_info = json.loads(_pinfo.read_text(encoding="utf-8")) if _pinfo.exists() else {}
        except Exception:
            _current_info = {}
        if _pname_to_save and _current_info.get("project_name") != _pname_to_save:
            _safe_num = re.sub(r'[<>:"/\\|?*\s]+', "_", project_number.strip()).strip("_")
            # Guard 1: folder name must match the sanitized project number.
            # If active_project_folder resolved to _APP_DIR (empty/invalid
            # project number), active_project_folder.name will not equal
            # _safe_num and the save is skipped.
            _name_matches = (active_project_folder != _APP_DIR and
                             active_project_folder.name == _safe_num)
            # Guard 2: never recreate a folder that was deleted this session.
            # _recently_deleted is permanent until New Project is clicked.
            _deleted_set = st.session_state.get("_recently_deleted", set())
            _recently_deleted_block = _safe_num in _deleted_set
            if not _name_matches:
                print(f"[SAVE BLOCKED] project_info.json save for {project_number!r} — folder name mismatch (active={active_project_folder.name!r}, safe_num={_safe_num!r})")
            elif _recently_deleted_block:
                print(f"[SAVE BLOCKED] project_info.json save for {project_number!r} — in _recently_deleted")
            else:
                print(f"[SAVE] SAVING project_info.json for: {project_number!r}")
                try:
                    # Intentional folder creation: user has entered both project
                    # number and project name, so register the project on disk.
                    active_project_folder.mkdir(parents=True, exist_ok=True)
                    _pinfo.write_text(
                        json.dumps({"project_number": project_number.strip(), "project_name": _pname_to_save}),
                        encoding="utf-8",
                    )
                except OSError:
                    pass

    # ── Load Existing Project ────────────────────────────────────────────────
    _existing = _list_projects()
    if _existing:
        _load_choice = st.selectbox(
            "Load Existing Project",
            options=["— select —"] + _existing,
            key="load_project_select",
        )
        if _load_choice and _load_choice != "— select —":
            _load_folder = PROJECTS_DIR / _load_choice
            _pinfo_path = _load_folder / "project_info.json"
            if st.button("Load", key="btn_load_project"):
                try:
                    _info = json.loads(_pinfo_path.read_text(encoding="utf-8")) if _pinfo_path.exists() else {}
                except Exception:
                    _info = {}
                # Use staging keys — writing directly to text_input widget-bound
                # keys after they've been rendered causes StreamlitAPIException.
                # The staging keys are consumed at the top of the next run.
                st.session_state["_staged_project_number"] = _info.get("project_number", _load_choice)
                st.session_state["_staged_project_name"] = _info.get("project_name", "")
                # Clear scope/drawing state so fresh restore happens on next run
                for _k in ["extracted_scope_items_text", "scope_summary_ready", "drawing_index_ready"]:
                    st.session_state.pop(_k, None)
                _restore_project_files(_load_folder)
                st.rerun()

    # ── New Project ──────────────────────────────────────────────────────────
    if st.button("New Project", key="btn_new_project"):
        # Only clear non-widget keys immediately; widget-bound keys are cleared
        # via the _reset_project flag at the top of the next run.
        for _k in [
            "extracted_scope_items_text", "scope_summary_ready",
            "drawing_index_ready", "appendix_b_scope_items",
            "division_ref", "appendix_b_intro", "file_labels",
            "_show_delete_confirm",
        ]:
            st.session_state.pop(_k, None)
        st.session_state["_reset_project"] = True
        st.rerun()

    st.divider()

    # ── Entity / Project type ────────────────────────────────────────────────
    st.subheader("Entity name")
    selected_entity = _exclusive_checkboxes("entity", ENTITY_OPTIONS)

    st.subheader("Project type")
    selected_project_type = _exclusive_checkboxes("ptype", PROJECT_TYPE_OPTIONS)

    st.divider()

    # ── Trade / Appendix B fields ────────────────────────────────────────────
    trade_or_division = st.selectbox(
        "Trade or Division",
        options=DIVISION_TABS,
        key="trade_or_division",
    )
    subcontractor_name = st.text_input(
        "Subcontractor Name",
        key="subcontractor_name",
        placeholder="Leave blank until contract award",
    )

    st.divider()

    # ── Estimator notes ──────────────────────────────────────────────────────
    st.subheader("Project Level Notes")
    estimator_notes = st.text_area(
        "Notes that apply to all trades for this project",
        height=120,
        placeholder="Phase constraints, owner-supplied items, global exclusions…",
        label_visibility="collapsed",
    )

    st.subheader("Division Level Notes")
    division_notes_trade = st.text_area(
        "Division Notes for this Trade",
        height=100,
        placeholder="Trade-specific scope instructions, exclusions, or clarifications.",
        label_visibility="collapsed",
    )

    # ── Delete Project ────────────────────────────────────────────────────────
    _proj_folder_exists = (
        bool(project_number)
        and (PROJECTS_DIR / re.sub(r'[<>:"/\\|?*\s]+', "_", project_number.strip()).strip("_")).exists()
    )
    if _proj_folder_exists:
        st.divider()
        if st.button(
            "🗑 Delete Project",
            key="btn_delete_project",
            type="primary",
            use_container_width=True,
        ):
            st.session_state["_show_delete_confirm"] = True

        if st.session_state.get("_show_delete_confirm"):
            st.warning(
                f"**Permanently delete project {project_number}?**\n\n"
                f"This will delete the project folder and all associated files. "
                f"This cannot be undone."
            )
            _confirm_text = st.text_input(
                'Type **DELETE** to confirm',
                key="_confirm_delete_input",
                placeholder="DELETE",
            )
            _del_col1, _del_col2 = st.columns(2)
            with _del_col1:
                if st.button(
                    "Confirm Delete",
                    key="btn_confirm_delete",
                    disabled=(_confirm_text.strip() != "DELETE"),
                    type="primary",
                    use_container_width=True,
                ):
                    print(f"[CONFIRM_DELETE] Button handler fired")
                    print(f"[CONFIRM_DELETE] active_project_folder = {active_project_folder!r}")
                    print(f"[CONFIRM_DELETE] Folder exists at click time: {active_project_folder.exists()}")
                    st.session_state["_delete_project_folder"] = str(active_project_folder)
                    print(f"[CONFIRM_DELETE] Stored _delete_project_folder = {st.session_state['_delete_project_folder']!r}")
                    st.session_state["_reset_project"] = True
                    st.session_state["_show_delete_confirm"] = False
                    st.rerun()
            with _del_col2:
                if st.button(
                    "Cancel",
                    key="btn_cancel_delete",
                    use_container_width=True,
                ):
                    st.session_state["_show_delete_confirm"] = False
                    st.rerun()


# ═══════════════════════════════════════════════════════════════════════════
# MAIN PANEL — File upload with type labeling
# ═══════════════════════════════════════════════════════════════════════════
st.subheader("Documents")

uploaded_pdfs = st.file_uploader(
    "Upload project PDFs",
    type=["pdf"],
    accept_multiple_files=True,
    label_visibility="collapsed",
)

# Per-file type labels stored in session state so they survive reruns
if "file_labels" not in st.session_state:
    st.session_state["file_labels"] = {}

drawings_files: list = []
spec_files: list = []

if uploaded_pdfs:
    st.caption("Label each file so it is routed to the correct pipeline:")
    for _uf in uploaded_pdfs:
        _col_name, _col_type = st.columns([3, 2])
        with _col_name:
            st.write(f"📄 {_uf.name}")
        with _col_type:
            _default_idx = 0
            _prev = st.session_state["file_labels"].get(_uf.name)
            if _prev and _prev in FILE_TYPE_OPTIONS:
                _default_idx = FILE_TYPE_OPTIONS.index(_prev)
            _label = st.selectbox(
                "Type",
                options=FILE_TYPE_OPTIONS,
                index=_default_idx,
                key=f"file_type_{_uf.name}",
                label_visibility="collapsed",
            )
            st.session_state["file_labels"][_uf.name] = _label
        if _label == "Drawings":
            drawings_files.append(_uf)
        elif _label == "Specifications":
            spec_files.append(_uf)


# ═══════════════════════════════════════════════════════════════════════════
# ACTIONS
# ═══════════════════════════════════════════════════════════════════════════
st.divider()

_bcol1, _bcol2, _bcol3 = st.columns(3)
with _bcol1:
    generate_scope_clicked = st.button("Generate Scope", use_container_width=True)
with _bcol2:
    appendix_b_clicked = st.button("Generate Appendix B", use_container_width=True)
with _bcol3:
    populate_car_clicked = st.button("Populate CAR", use_container_width=True)

# ── Generate Scope ────────────────────────────────────────────────────────────
if generate_scope_clicked:
    if not project_number:
        st.warning("Enter a Project Number in the sidebar before generating scope.")
    elif not drawings_files and not spec_files:
        st.warning(
            "No files labeled as Drawings or Specifications. "
            "Label at least one uploaded PDF before running Generate Scope."
        )
    else:
        _pf = active_project_folder

        # ── Index all Drawings files ──────────────────────────────────────
        if drawings_files:
            st.subheader("Drawing Indexing")
            if len(drawings_files) == 1:
                st.info(f"Indexing drawings: {drawings_files[0].name}")
            else:
                st.info(f"Indexing {len(drawings_files)} drawings file(s): " +
                        ", ".join(f.name for f in drawings_files))
            _combined_index: list[dict] = []
            for _df in drawings_files:
                try:
                    index_drawings(
                        _df.getvalue(),
                        project_number,
                        project_name,
                        project_folder=_pf,
                    )
                except Exception as exc:
                    st.error(f"Could not index {_df.name}: {exc}")
        else:
            st.warning("No files labeled 'Drawings' — drawing index skipped.")

        # ── Extract all Specifications files ──────────────────────────────
        if spec_files:
            st.subheader("Specification Extraction")
            combined_chunks: list[str] = []
            for _sf in spec_files:
                st.info(f"Extracting text: {_sf.name}")
                combined_chunks.append(f"--- {_sf.name} ---\n")
                try:
                    combined_chunks.append(extract_pdf_text(_sf.getvalue()))
                except Exception as exc:
                    st.error(f"Could not extract text from {_sf.name}: {exc}")
                combined_chunks.append("\n\n")
            if combined_chunks:
                try:
                    parse_spec_division("".join(combined_chunks), project_folder=_pf)
                    scope_items_text = st.session_state.get("extracted_scope_items_text")
                    if scope_items_text:
                        docx_bytes = _scope_items_to_docx_bytes(scope_items_text)
                        st.download_button(
                            label="Download Scope Summary (.docx)",
                            data=docx_bytes,
                            file_name=_spaced_download_filename(
                                [_project_field(project_number), _project_field(project_name), "Scope Summary"],
                                "docx",
                            ),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                except Exception as exc:
                    st.error(f"Specification extraction failed: {exc}")
        else:
            st.warning("No files labeled 'Specifications' — spec extraction skipped.")

# ── Generate Appendix B ───────────────────────────────────────────────────────
if appendix_b_clicked:
    _pf = active_project_folder
    _missing = []
    if not _has_scope_summary_output():
        _missing.append("Generate Scope (Specifications)")
    if not _has_drawing_index_file(_pf):
        _missing.append("Generate Scope (Drawings)")
    if _missing:
        st.warning("Run " + " and ".join(_missing) + " before generating Appendix B.")
    elif not trade_or_division:
        st.warning("Select a Trade or Division in the sidebar.")
    else:
        try:
            generate_appendix_b(
                estimator_notes,
                division_notes_trade,
                trade_or_division.strip(),
                project_folder=_pf,
            )
            appendix_text = st.session_state.get("appendix_b_scope_items")
            if appendix_text:
                _sub = (subcontractor_name or "").strip()
                _filename_trade = _sub if _sub else trade_or_division.strip()
                ab_docx = _appendix_b_word_bytes(
                    project_number,
                    trade_or_division.strip(),
                    st.session_state.get("division_ref") or "",
                    appendix_text,
                    selected_entity,
                    intro_text=st.session_state.get("appendix_b_intro") or "",
                    subcontractor_name=subcontractor_name or "",
                )
                st.download_button(
                    label="Download Appendix B (.docx)",
                    data=ab_docx,
                    file_name=_spaced_download_filename(
                        [
                            _project_field(project_number),
                            _project_field(project_name),
                            "Appendix B",
                            _project_field(_filename_trade),
                        ],
                        "docx",
                    ),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except Exception as exc:
            st.error(f"Could not generate Appendix B: {exc}")

# ── Populate CAR ─────────────────────────────────────────────────────────────
if populate_car_clicked:
    st.info("Coming soon — upload trade quotes to populate CAR.")
